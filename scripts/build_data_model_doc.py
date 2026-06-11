from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Yapsu-Data-Model-and-Relationships.docx"
ASSET_DIR = ROOT / "tmp" / "data-model-doc-assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "16324F"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "E7F4ED"
LIGHT_GOLD = "FFF4D6"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "202124"
GRID = "CCD3DA"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_callout(doc, title, text, fill=LIGHT_BLUE, title_color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=fill, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run_font(r2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=font_size, bold=True, color=NAVY)
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_run_font(r, size=font_size)
    return table


def font(size=16, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def draw_box(draw, xy, title, lines, fill, outline="#667085"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=12, fill=fill, outline=outline, width=2)
    draw.text((x1 + 14, y1 + 10), title, font=font(17, True), fill="#16324F")
    y = y1 + 38
    for line in lines:
        draw.text((x1 + 14, y), line, font=font(12), fill="#344054")
        y += 19


def arrow(draw, start, end, label=None):
    draw.line([start, end], fill="#475467", width=3)
    x2, y2 = end
    x1, y1 = start
    angle_horizontal = abs(x2 - x1) >= abs(y2 - y1)
    if angle_horizontal:
        direction = 1 if x2 > x1 else -1
        pts = [(x2, y2), (x2 - 10 * direction, y2 - 6), (x2 - 10 * direction, y2 + 6)]
    else:
        direction = 1 if y2 > y1 else -1
        pts = [(x2, y2), (x2 - 6, y2 - 10 * direction), (x2 + 6, y2 - 10 * direction)]
    draw.polygon(pts, fill="#475467")
    if label:
        mx, my = (x1 + x2) // 2, (y1 + y2) // 2
        draw.rounded_rectangle((mx - 32, my - 10, mx + 32, my + 10), 4, fill="#FFFFFF")
        draw.text((mx - 26, my - 7), label, font=font(10, True), fill="#475467")


def create_domain_diagram(path):
    image = Image.new("RGB", (1600, 1000), "#F7F9FC")
    draw = ImageDraw.Draw(image)
    draw.text((50, 25), "Yapsu - Proposed Domain Relationship Overview", font=font(28, True), fill="#16324F")

    boxes = {
        "lang": (60, 110, 345, 275),
        "curr": (410, 110, 700, 275),
        "lesson": (765, 110, 1055, 275),
        "content": (1115, 110, 1535, 275),
        "audio": (1115, 345, 1535, 530),
        "drill": (765, 610, 1055, 800),
        "role": (410, 610, 700, 800),
        "loc": (60, 610, 345, 800),
        "qa": (1115, 610, 1535, 800),
    }
    draw_box(draw, boxes["lang"], "Language domain", ["languages", "language_pairs", "pair_locales"], "#E7F4ED")
    draw_box(draw, boxes["curr"], "Curriculum", ["curricula", "curriculum_imports", "import_rows"], "#E8EEF5")
    draw_box(draw, boxes["lesson"], "Lesson", ["lessons", "lesson_locales", "publish readiness"], "#FFF4D6")
    draw_box(draw, boxes["content"], "Content", ["content_items", "content_revisions", "A / V / S / G types"], "#E8EEF5")
    draw_box(draw, boxes["audio"], "Audio", ["audio_assets", "audio_revisions", "source revision linkage"], "#FCE8E6")
    draw_box(draw, boxes["drill"], "Drill", ["drill_assignments", "drill_configs", "blank token ranges"], "#E7F4ED")
    draw_box(draw, boxes["role"], "Roleplay", ["roleplays", "roleplay_goals", "mobile description"], "#FFF4D6")
    draw_box(draw, boxes["loc"], "Localization", ["content_localizations", "roleplay localizations", "locale coverage"], "#F3E8FF")
    draw_box(draw, boxes["qa"], "Quality control", ["ai_qa_results", "human_qa_reviews", "approval history"], "#FCE8E6")

    arrow(draw, (345, 190), (410, 190), "1:N")
    arrow(draw, (700, 190), (765, 190), "1:N")
    arrow(draw, (1055, 190), (1115, 190), "1:N")
    arrow(draw, (1325, 275), (1325, 345), "1:N")
    arrow(draw, (1115, 700), (1055, 700), "N:1")
    arrow(draw, (765, 690), (700, 690), "N:1")
    arrow(draw, (410, 690), (345, 690), "1:N")
    arrow(draw, (1325, 530), (1325, 610), "1:N")
    arrow(draw, (910, 275), (910, 610), "1:N")
    arrow(draw, (1170, 275), (330, 610), "1:N")

    image.save(path)


def create_dependency_diagram(path):
    image = Image.new("RGB", (1600, 780), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.text((50, 25), "Change Propagation: What happens when source content changes?", font=font(27, True), fill="#16324F")

    stages = [
        ((70, 160, 330, 310), "Content Revision", ["Operator edits script,", "vocab or sentence", "revision N -> N+1"], "#E8EEF5"),
        ((395, 160, 655, 310), "Audio Status", ["Existing audio becomes", "OUTDATED because it", "points to revision N"], "#FCE8E6"),
        ((720, 160, 980, 310), "QA Status", ["AI/Human QA results", "remain audit history", "but are no longer valid"], "#FFF4D6"),
        ((1045, 160, 1305, 310), "Downstream Review", ["Drill becomes Needs Review", "Roleplay flagged only when", "it references changed item"], "#F3E8FF"),
        ((565, 470, 1035, 650), "Lesson Publish Readiness", ["Blocked until required localization, audio and QA are valid.", "After regeneration and approval, readiness can return to Ready."], "#E7F4ED"),
    ]
    for xy, title, lines, fill in stages:
        draw_box(draw, xy, title, lines, fill)

    arrow(draw, (330, 235), (395, 235), "event")
    arrow(draw, (655, 235), (720, 235), "invalid")
    arrow(draw, (980, 235), (1045, 235), "flag")
    arrow(draw, (850, 310), (800, 470), "gate")
    arrow(draw, (1175, 310), (900, 470), "gate")
    image.save(path)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167


def add_cover(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("YAPSU · INTERNAL PRODUCT SPECIFICATION")
    set_run_font(r, size=8.5, color=MUTED, bold=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(56)
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(10)
    kr = kicker.add_run("CURRICULUM MANAGEMENT PLATFORM")
    set_run_font(kr, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    tr = title.add_run("Data Model &\nData Relationships")
    set_run_font(tr, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(26)
    sr = subtitle.add_run(
        "Proposed logical model for curriculum authoring, audio versioning, "
        "quality assurance, drills, roleplay and localization."
    )
    set_run_font(sr, size=13, color=MUTED)

    meta = [
        ("Document status", "Proposed MVP baseline"),
        ("Prepared for", "Yapsu product and engineering review"),
        ("Prepared date", "11 June 2026"),
        ("Source context", "README, UI review transcript and current prototype"),
    ]
    table = add_table(doc, ["Document", "Value"], meta, [2200, 7160], font_size=10)
    table.rows[0]._element.getparent().remove(table.rows[0]._element)

    doc.add_paragraph().paragraph_format.space_after = Pt(22)
    add_callout(
        doc,
        "Core design decision",
        "Stable business identities are separated from editable revisions. "
        "Audio and QA records always point to the exact content revision they were created from.",
        fill=LIGHT_GREEN,
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(36)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Internal use only · Version 1.0")
    set_run_font(r, size=9, color=MUTED, italic=True)
    doc.add_page_break()


def add_toc(doc):
    add_heading(doc, "Contents", 1)
    entries = [
        "1. Purpose and scope",
        "2. Modeling principles",
        "3. Domain relationship overview",
        "4. Entity catalog",
        "5. Detailed data dictionary",
        "6. Status and versioning rules",
        "7. Change propagation and dependency matrix",
        "8. Validation and business rules",
        "9. Use-case mapping",
        "10. Recommended implementation sequence",
        "Appendix A. Enumerations",
        "Appendix B. Open decisions",
    ]
    for entry in entries:
        add_bullet(doc, entry)
    add_callout(
        doc,
        "Document intent",
        "This is a logical data specification for product alignment. "
        "Physical PostgreSQL/Supabase DDL should be produced only after open decisions are confirmed.",
        fill=LIGHT_GOLD,
        title_color="7A5A00",
    )
    doc.add_page_break()


def build_document():
    diagram1 = ASSET_DIR / "domain-overview.png"
    diagram2 = ASSET_DIR / "change-propagation.png"
    create_domain_diagram(diagram1)
    create_dependency_diagram(diagram2)

    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_toc(doc)

    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("Yapsu Data Model · ")
    set_run_font(fr, size=8.5, color=MUTED)
    add_page_number(fp)

    add_heading(doc, "1. Purpose and scope", 1)
    add_body(
        doc,
        "This document defines the proposed logical data model for the Yapsu Curriculum Management Platform. "
        "It covers the complete authoring pipeline from curriculum import to content editing, audio production, "
        "AI and human QA, drill assignment, roleplay configuration, localization and release to the mobile app."
    )
    add_body(
        doc,
        "The model is designed to answer the main product questions raised during the UI review: which fields are editable, "
        "how completion and enablement are represented, how an audio file is tied to the correct script version, and which "
        "downstream records become invalid when a vocab, sentence or tutor script changes."
    )

    add_heading(doc, "1.1 In scope", 2)
    for item in [
        "Logical entities, keys and relationships.",
        "Versioning and audit history for curriculum content and audio.",
        "Localization coverage by language pair and lesson.",
        "AI QA and Human QA as separate stages.",
        "Exclusive Drill versus Extra Drill assignment.",
        "Roleplay description and goal mapping to the mobile app.",
        "Import traceability and publish-readiness rules.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.2 Out of scope", 2)
    for item in [
        "Authentication provider details and final Row Level Security policies.",
        "Gemini API request/response payloads and storage provider configuration.",
        "Final analytics, learner progress and mobile runtime data.",
        "Exact SQL indexes and performance tuning.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. Modeling principles", 1)
    principles = [
        ("Stable identity, versioned content", "A content item keeps one stable ID; each edit creates a new content revision."),
        ("Explicit source linkage", "Audio, QA and drill records reference a concrete source revision, never only free text."),
        ("No silent invalidation", "When source content changes, affected downstream records receive an explicit status."),
        ("Structured configuration", "Drill configuration stores typed fields or token ranges instead of freestyle text fragments."),
        ("Localization as first-class data", "Native-language coverage and enablement are modeled per locale and lesson."),
        ("Auditability", "Imports, revisions, QA decisions and publish transitions keep timestamps and actor IDs."),
        ("Mobile mapping before UI fields", "A CMS field exists only when it has a defined business purpose or mobile mapping."),
    ]
    add_table(doc, ["Principle", "Meaning"], principles, [2500, 6860], font_size=9.5)

    add_heading(doc, "3. Domain relationship overview", 1)
    add_body(doc, "Figure 1 shows the proposed domain boundaries and the primary relationship direction.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(diagram1), width=Inches(6.35))
    cp = doc.add_paragraph("Figure 1. Proposed logical domain relationship overview.")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.runs[0].italic = True
    cp.runs[0].font.size = Pt(9)
    cp.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    add_heading(doc, "3.1 Cardinality summary", 2)
    relations = [
        ("Language", "Language Pair", "1:N", "A language can participate in many learning/native pairs."),
        ("Language Pair", "Curriculum", "1:N", "A pair can have multiple curricula or curriculum versions."),
        ("Curriculum", "Lesson", "1:N", "A curriculum contains ordered lessons."),
        ("Lesson", "Content Item", "1:N", "A lesson owns Tutor, Vocab, Sentence and Grammar items."),
        ("Content Item", "Content Revision", "1:N", "Every edit produces a revision while identity remains stable."),
        ("Content Revision", "Audio Revision", "1:N", "Audio is generated/uploaded from one exact source revision."),
        ("Audio Revision", "AI QA Result", "1:N", "AI QA may be rerun and results are retained."),
        ("Audio Revision", "Human QA Review", "1:N", "Human decisions are auditable and may be superseded."),
        ("Content Item", "Drill Assignment", "0:1", "One source card may be unassigned, Drill or Extra Drill."),
        ("Lesson", "Roleplay", "0:1", "A lesson may have one roleplay configuration."),
        ("Roleplay", "Roleplay Goal", "1:N", "Goals are ordered and can be enabled independently."),
        ("Content Revision", "Localization", "1:N", "Each revision can have multiple locale overlays."),
    ]
    add_table(doc, ["Parent", "Child", "Cardinality", "Rule"], relations, [1650, 1850, 1000, 4860], font_size=8.7)

    add_heading(doc, "4. Entity catalog", 1)
    entities = [
        ("languages", "Master", "Supported languages and locale metadata.", "UC01"),
        ("language_pairs", "Master", "Learning language to native language configuration.", "UC01"),
        ("curricula", "Core", "Imported or authored curriculum container.", "UC01"),
        ("curriculum_imports", "Audit", "Upload batch, validation and import result.", "UC01"),
        ("lessons", "Core", "Ordered lesson metadata and release status.", "UC01"),
        ("lesson_locales", "Localization", "Locale availability, coverage and enablement per lesson.", "UC01/04"),
        ("content_items", "Core", "Stable identity for Tutor/Vocab/Sentence/Grammar cards.", "UC01"),
        ("content_revisions", "Versioned", "Editable content snapshots and revision metadata.", "UC01"),
        ("content_localizations", "Localization", "Native translations tied to a content revision.", "UC01/02/03"),
        ("audio_assets", "Media", "Stable audio identity for a content item.", "UC02"),
        ("audio_revisions", "Versioned", "Uploaded/generated files tied to source revision.", "UC02"),
        ("ai_qa_results", "QA", "Automated quality checks and details.", "UC02"),
        ("human_qa_reviews", "QA", "Human Pass/Fail decisions and reasons.", "UC02"),
        ("drill_assignments", "Interactive", "Exclusive Drill/Extra Drill placement.", "UC03"),
        ("drill_configs", "Interactive", "Typed drill configuration.", "UC03"),
        ("roleplays", "Interactive", "Mobile-facing roleplay description and status.", "UC04"),
        ("roleplay_goals", "Interactive", "Ordered success goals and native descriptions.", "UC04"),
        ("publish_checks", "Derived/Audit", "Readiness checks and blocking reasons.", "All"),
    ]
    add_table(doc, ["Entity", "Category", "Purpose", "UC"], entities, [2100, 1450, 4710, 1100], font_size=8.6)

    add_heading(doc, "5. Detailed data dictionary", 1)
    add_body(doc, "The following fields define the MVP baseline. UUID is recommended for primary keys; timestamps use UTC.")

    dictionaries = {
        "5.1 languages": [
            ("id", "uuid", "PK", "Stable language ID."),
            ("code", "varchar(10)", "UK, required", "BCP-47/ISO-like code, e.g. zh, vi, ja."),
            ("name", "varchar(100)", "required", "Display name."),
            ("native_name", "varchar(100)", "required", "Native display name."),
            ("is_active", "boolean", "required", "Whether language can be selected."),
        ],
        "5.2 language_pairs": [
            ("id", "uuid", "PK", "Stable pair ID."),
            ("learning_language_id", "uuid", "FK languages", "Language being learned."),
            ("native_language_id", "uuid", "FK languages", "Learner UI/support language."),
            ("status", "enum", "required", "draft, active, disabled."),
            ("created_at / updated_at", "timestamptz", "required", "Audit timestamps."),
        ],
        "5.3 curricula": [
            ("id", "uuid", "PK", "Curriculum identity."),
            ("language_pair_id", "uuid", "FK language_pairs", "Pair served by curriculum."),
            ("code", "varchar(50)", "UK, required", "Stable business code."),
            ("title", "varchar(255)", "required", "Curriculum name."),
            ("version_label", "varchar(50)", "optional", "Human-readable release label."),
            ("status", "enum", "required", "draft, validating, ready, published, archived."),
            ("current_import_id", "uuid", "FK curriculum_imports", "Latest accepted import batch."),
        ],
        "5.4 curriculum_imports": [
            ("id", "uuid", "PK", "Import batch."),
            ("curriculum_id", "uuid", "FK curricula", "Target curriculum."),
            ("file_name / file_url", "text", "required", "Source Excel reference."),
            ("status", "enum", "required", "uploaded, validating, valid, invalid, imported, cancelled."),
            ("valid_row_count / error_row_count", "integer", "required", "Preview summary."),
            ("validation_report", "jsonb", "required", "Structured row/field errors."),
            ("created_by / created_at", "uuid / timestamptz", "required", "Audit data."),
        ],
        "5.5 lessons": [
            ("id", "uuid", "PK", "Lesson identity."),
            ("curriculum_id", "uuid", "FK curricula", "Owning curriculum."),
            ("lesson_code", "varchar(50)", "UK within curriculum", "Example CN_L101."),
            ("level / order_index", "integer", "required", "Curriculum ordering."),
            ("title", "varchar(255)", "required", "Inherited by Roleplay; edited here."),
            ("description", "text", "optional", "Lesson summary."),
            ("status", "enum", "required", "draft, incomplete, ready, published, disabled, needs_review."),
            ("is_enabled", "boolean", "required", "Operator-controlled availability after validation."),
            ("published_revision_at", "timestamptz", "optional", "Last successful publish."),
        ],
        "5.6 lesson_locales": [
            ("id", "uuid", "PK", "Locale coverage record."),
            ("lesson_id", "uuid", "FK lessons", "Lesson."),
            ("locale_language_id", "uuid", "FK languages", "Native locale."),
            ("is_enabled", "boolean", "required", "Locale available for this lesson."),
            ("coverage_status", "enum", "required", "missing, partial, complete, needs_review."),
            ("coverage_percent", "numeric(5,2)", "derived/cache", "UI progress indicator."),
        ],
        "5.7 content_items": [
            ("id", "uuid", "PK", "Stable card identity."),
            ("lesson_id", "uuid", "FK lessons", "Owning lesson."),
            ("code", "varchar(80)", "UK, required", "Example CN_L101_V6."),
            ("content_type", "enum", "required", "tutor, vocab, sentence, grammar."),
            ("order_index", "integer", "required", "Display/teaching sequence."),
            ("is_enabled", "boolean", "required", "Whether item participates in release."),
            ("current_revision_id", "uuid", "FK content_revisions", "Currently selected revision."),
            ("status", "enum", "required", "incomplete, ready, disabled, needs_review."),
        ],
        "5.8 content_revisions": [
            ("id", "uuid", "PK", "Immutable revision snapshot."),
            ("content_item_id", "uuid", "FK content_items", "Stable item."),
            ("revision_no", "integer", "required", "Monotonic revision number."),
            ("script_text", "text", "conditional", "Target-language text or Tutor script."),
            ("reading", "text", "optional", "Pinyin/kana/romanization."),
            ("meaning_en", "text", "optional", "Canonical English meaning/explanation."),
            ("content_hash", "varchar(64)", "required", "Detects exact content changes."),
            ("change_note", "text", "optional", "Reason for revision."),
            ("created_by / created_at", "uuid / timestamptz", "required", "Audit metadata."),
        ],
        "5.9 content_localizations": [
            ("id", "uuid", "PK", "Localization record."),
            ("content_revision_id", "uuid", "FK content_revisions", "Exact source revision."),
            ("locale_language_id", "uuid", "FK languages", "Target locale."),
            ("localized_text", "text", "required", "Native translation/description."),
            ("status", "enum", "required", "draft, reviewed, approved, outdated."),
            ("reviewed_by / reviewed_at", "uuid / timestamptz", "optional", "Localization QA."),
        ],
        "5.10 audio_assets and audio_revisions": [
            ("audio_assets.id", "uuid", "PK", "Stable audio identity per content item/voice role."),
            ("audio_assets.content_item_id", "uuid", "FK content_items", "Source card."),
            ("audio_revisions.id", "uuid", "PK", "Audio file revision."),
            ("audio_revisions.audio_asset_id", "uuid", "FK audio_assets", "Stable audio identity."),
            ("audio_revisions.source_revision_id", "uuid", "FK content_revisions", "Exact script used."),
            ("source_type", "enum", "required", "uploaded or generated."),
            ("file_url / mime_type / duration_ms", "text", "required after ready", "Media metadata."),
            ("provider / voice_config", "text / jsonb", "optional", "Gemini/provider settings."),
            ("status", "enum", "required", "missing, generating, ready, failed, outdated, approved."),
            ("generation_message", "text", "optional", "Provider error or status detail."),
        ],
        "5.11 ai_qa_results and human_qa_reviews": [
            ("audio_revision_id", "uuid", "FK audio_revisions", "Audio under review."),
            ("result", "enum", "required", "pending, passed, failed, error."),
            ("checks", "jsonb", "AI QA", "Pronunciation, silence, clipping, alignment, etc."),
            ("score", "numeric", "AI QA optional", "Normalized quality score."),
            ("reason", "text", "Human QA required on fail", "Reviewer explanation."),
            ("reviewed_by", "uuid", "Human QA", "Reviewer ID."),
            ("created_at", "timestamptz", "required", "Decision timestamp."),
            ("is_current", "boolean", "required", "Latest applicable result for revision."),
        ],
        "5.12 drill_assignments and drill_configs": [
            ("drill_assignments.id", "uuid", "PK", "Assignment identity."),
            ("lesson_id", "uuid", "FK lessons", "Owning lesson."),
            ("source_content_item_id", "uuid", "FK content_items, UK", "One assignment per source card."),
            ("placement", "enum", "required", "drill or extra_drill."),
            ("drill_type", "enum", "required", "listen_repeat, fill_blank, sentence_order."),
            ("status", "enum", "required", "draft, ready, needs_review, disabled."),
            ("source_revision_id", "uuid", "FK content_revisions", "Revision last reviewed."),
            ("config", "jsonb", "typed by drill_type", "Type-specific structured data."),
            ("blank_start / blank_end", "integer", "fill_blank", "Contiguous token range."),
        ],
        "5.13 roleplays and roleplay_goals": [
            ("roleplays.id", "uuid", "PK", "Roleplay identity."),
            ("lesson_id", "uuid", "FK lessons, UK", "One roleplay per lesson."),
            ("description", "text", "required", "Mobile-facing context/description."),
            ("status", "enum", "required", "draft, ready, disabled, needs_review."),
            ("roleplay_goals.id", "uuid", "PK", "Goal identity."),
            ("roleplay_id", "uuid", "FK roleplays", "Parent roleplay."),
            ("order_index", "integer", "required", "Display/evaluation order."),
            ("goal_en", "text", "required", "Canonical success criterion."),
            ("description_native", "text", "required per enabled locale", "Learner-facing native description."),
            ("is_enabled", "boolean", "required", "Whether goal is active."),
        ],
    }

    for title, rows in dictionaries.items():
        add_heading(doc, title, 2)
        add_table(doc, ["Field", "Type", "Constraint", "Purpose"], rows, [2200, 1550, 2100, 3510], font_size=8.5)

    add_heading(doc, "6. Status and versioning rules", 1)
    add_heading(doc, "6.1 Content revision rule", 2)
    add_body(
        doc,
        "Editing script_text, reading, meaning_en or another release-significant field creates a new immutable content revision. "
        "The content item points current_revision_id to the new revision. Previous revisions remain available for audit and rollback."
    )
    add_heading(doc, "6.2 Audio synchronization rule", 2)
    add_body(
        doc,
        "An audio revision is synchronized only when audio_revisions.source_revision_id equals content_items.current_revision_id. "
        "If they differ, the UI must display Outdated and offer Upload New Audio or Regenerate."
    )
    add_heading(doc, "6.3 QA applicability rule", 2)
    add_body(
        doc,
        "AI and Human QA results apply to one audio revision only. A new audio revision starts with AI QA Pending and Human QA Pending. "
        "Old decisions remain historical but cannot approve the new file."
    )
    add_heading(doc, "6.4 Publish readiness", 2)
    checks = [
        ("Required content", "All enabled cards have a complete current revision.", "Blocking"),
        ("Localization", "All enabled lesson locales meet required coverage.", "Blocking"),
        ("Audio", "Required cards have synchronized, ready audio.", "Blocking"),
        ("AI QA", "Latest required AI QA result is Passed.", "Configurable blocking"),
        ("Human QA", "Latest Human QA result is Passed.", "Blocking"),
        ("Drill", "Assignments reference the current reviewed content revision.", "Blocking when enabled"),
        ("Roleplay", "Description exists and at least one enabled valid goal exists.", "Blocking when roleplay enabled"),
    ]
    add_table(doc, ["Check", "Rule", "Severity"], checks, [1900, 5960, 1500], font_size=9)

    add_heading(doc, "7. Change propagation and dependency matrix", 1)
    add_body(doc, "Figure 2 describes the required system response when source content changes.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(diagram2), width=Inches(6.35))
    cp = doc.add_paragraph("Figure 2. Source edit propagation and release gate.")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.runs[0].italic = True
    cp.runs[0].font.size = Pt(9)
    cp.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    matrix = [
        ("Lesson title", "Roleplay header/mobile lesson title", "No audio impact", "Roleplay preview refresh", "Revalidate mobile mapping"),
        ("Tutor script", "Tutor content revision", "Audio Outdated", "Reset applicable QA", "Lesson Needs Review"),
        ("Vocab script/reading", "Vocab revision", "Audio Outdated", "Drill Needs Review", "Localization may become Outdated"),
        ("Sentence script/reading", "Sentence revision", "Audio Outdated", "Fill Blank/Sentence Order Needs Review", "Token ranges must be revalidated"),
        ("Meaning/translation", "Localization/current revision", "No TTS impact unless spoken", "Drill native meaning refresh", "Locale coverage may become partial"),
        ("Disable content item", "Release composition", "Audio retained as history", "Disable dependent assignment", "Recalculate readiness"),
        ("Audio file", "Audio revision", "New current audio", "AI/Human QA Pending", "Cannot publish until approved"),
        ("Language locale enable", "lesson_locales", "No direct impact", "Requires localization coverage", "Block locale publish if incomplete"),
    ]
    add_table(
        doc,
        ["Changed field", "Primary effect", "Audio", "Downstream", "Release impact"],
        matrix,
        [1450, 1900, 1500, 2610, 1900],
        font_size=7.9,
    )

    add_heading(doc, "8. Validation and business rules", 1)
    rules = [
        ("BR-01", "A language pair cannot use the same learning and native language.", "Database constraint"),
        ("BR-02", "lesson_code is unique within a curriculum.", "Unique constraint"),
        ("BR-03", "content item code is unique and its type prefix must match content_type.", "Validation + unique constraint"),
        ("BR-04", "An enabled content item must have a complete current revision.", "Application + publish check"),
        ("BR-05", "Audio Approved requires synchronized source revision and current Human QA Passed.", "State transition guard"),
        ("BR-06", "A source content item can have at most one active Drill Assignment.", "Unique constraint"),
        ("BR-07", "Drill and Extra Drill are mutually exclusive placements.", "Enum + unique constraint"),
        ("BR-08", "Fill Blank selection must be a non-empty contiguous token range.", "Typed config validation"),
        ("BR-09", "Roleplay requires at least one enabled, complete goal when enabled.", "Publish check"),
        ("BR-10", "Lesson locale cannot be enabled when required localization coverage is incomplete.", "State transition guard"),
        ("BR-11", "Hard/master fields are selected from controlled values, not free edited.", "UI + FK constraint"),
        ("BR-12", "Delete operations on referenced records default to soft delete or restrict.", "FK/repository policy"),
    ]
    add_table(doc, ["ID", "Rule", "Enforcement"], rules, [900, 6260, 2200], font_size=8.7)

    add_heading(doc, "9. Use-case mapping", 1)
    uc_rows = [
        ("UC01 Curriculum Editor", "curricula, imports, lessons, content_items, content_revisions, lesson_locales", "Validated curriculum with explicit readiness and locale coverage."),
        ("UC02 Audio QA", "audio_assets, audio_revisions, ai_qa_results, human_qa_reviews", "Approved audio synchronized to the current script revision."),
        ("UC03 Drill Editor", "drill_assignments, drill_configs, content_items/revisions", "Typed Drill/Extra Drill configurations without duplicate assignment."),
        ("UC04 Roleplay Editor", "roleplays, roleplay_goals, lessons, lesson_locales", "Mobile description plus ordered active success goals."),
    ]
    add_table(doc, ["Use Case", "Primary entities", "Output"], uc_rows, [1900, 4260, 3200], font_size=8.8)

    add_heading(doc, "10. Recommended implementation sequence", 1)
    sequence = [
        ("1", "Master and core structure", "languages, language_pairs, curricula, lessons"),
        ("2", "Versioned curriculum content", "content_items, content_revisions, content_localizations"),
        ("3", "Import workflow", "curriculum_imports, validation report and preview mapping"),
        ("4", "Audio and QA", "audio assets/revisions, AI QA, Human QA, stale detection"),
        ("5", "Interactive modules", "drill assignments/configs, roleplay/goals"),
        ("6", "Release controls", "publish checks, status transitions and mobile export view"),
        ("7", "Security and operations", "Supabase RLS, audit events, storage policies and jobs"),
    ]
    add_table(doc, ["Phase", "Scope", "Tables/Capability"], sequence, [900, 2750, 5710], font_size=9)

    add_heading(doc, "Appendix A. Enumerations", 1)
    enums = [
        ("content_type", "tutor, vocab, sentence, grammar"),
        ("lesson_status", "draft, incomplete, ready, published, disabled, needs_review"),
        ("coverage_status", "missing, partial, complete, needs_review"),
        ("audio_status", "missing, generating, ready, failed, outdated, approved"),
        ("qa_result", "pending, passed, failed, error"),
        ("drill_placement", "drill, extra_drill"),
        ("drill_type", "listen_repeat, fill_blank, sentence_order"),
        ("record_status", "draft, ready, disabled, needs_review"),
        ("import_status", "uploaded, validating, valid, invalid, imported, cancelled"),
    ]
    add_table(doc, ["Enum", "Values"], enums, [2600, 6760], font_size=9)

    add_heading(doc, "Appendix B. Open decisions", 1)
    decisions = [
        ("AI QA blocking?", "Confirm whether AI QA Passed is mandatory for publish or advisory before Human QA."),
        ("Audio granularity", "Confirm which content types require audio: Tutor only or Tutor/Vocab/Sentence/Grammar."),
        ("Canonical English", "Confirm whether English is a fixed canonical meaning or another localization."),
        ("Roleplay prompt ownership", "Confirm that operators edit only mobile description/goals, not the internal system prompt."),
        ("Multiple curricula", "Confirm whether one language pair can have multiple active curricula or only archived versions."),
        ("Import merge behavior", "Define replace, merge and duplicate handling for an existing lesson/card code."),
        ("Deletion policy", "Confirm soft-delete retention period and whether published records can ever be hard deleted."),
        ("Mobile export contract", "Define final JSON/API schema and which statuses are visible to mobile."),
    ]
    add_table(doc, ["Decision", "Question to confirm"], decisions, [2650, 6710], font_size=9)

    add_callout(
        doc,
        "Approval recommendation",
        "Review and confirm the open decisions first. After approval, convert this logical model into Supabase SQL migrations, "
        "RLS policies and typed TypeScript models. UI implementation should use the same status and dependency rules.",
        fill=LIGHT_GREEN,
    )

    doc.core_properties.title = "Yapsu Data Model and Data Relationships"
    doc.core_properties.subject = "Curriculum Management Platform logical data specification"
    doc.core_properties.author = "Yapsu"
    doc.core_properties.keywords = "Yapsu, curriculum, data model, relationships, audio QA, drill, roleplay"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
